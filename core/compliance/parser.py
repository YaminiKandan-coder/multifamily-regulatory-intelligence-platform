from __future__ import annotations
import re
from dataclasses import dataclass, field
from typing import Optional


@dataclass
class Clause:
    number: str
    title: str
    content: str


@dataclass
class ParsedDocument:
    clauses: list[Clause] = field(default_factory=list)
    raw_text: str = ""


def parse_pdf(file_bytes: bytes) -> str:
    import PyPDF2
    import io
    reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def parse_docx(file_bytes: bytes) -> str:
    import docx
    import io
    doc = docx.Document(io.BytesIO(file_bytes))
    return "\n".join(p.text for p in doc.paragraphs)


def parse_document(file_bytes: bytes, filename: str) -> str:
    ext = filename.lower().rsplit(".", 1)[-1]
    if ext == "pdf":
        return parse_pdf(file_bytes)
    if ext in ("docx", "doc"):
        return parse_docx(file_bytes)
    return file_bytes.decode("utf-8", errors="replace")


# Numbered clause pattern: "1.", "1.1.", "Section 1", "Article I"
_NUMBERED_RE = re.compile(
    r"^(\d+\.[\d\.]*|Section\s+\d+[\.\d]*|Article\s+[IVXLCDM\d]+)\s*[:\-–]?\s*(.{0,80})",
    re.MULTILINE | re.IGNORECASE,
)
# ALL-CAPS heading fallback
_CAPS_RE = re.compile(r"^([A-Z][A-Z\s]{4,50})$", re.MULTILINE)


def extract_clauses(text: str) -> ParsedDocument:
    matches = list(_NUMBERED_RE.finditer(text))
    if len(matches) >= 2:
        clauses = _extract_by_matches(text, matches)
    else:
        fallback_matches = list(_CAPS_RE.finditer(text))
        clauses = _extract_by_matches(text, fallback_matches) if fallback_matches else [
            Clause(number="1", title="Full Document", content=text.strip())
        ]
    return ParsedDocument(clauses=clauses, raw_text=text)


def _extract_by_matches(text: str, matches: list) -> list[Clause]:
    clauses: list[Clause] = []
    for i, m in enumerate(matches):
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        number = m.group(1).strip() if m.lastindex and m.lastindex >= 1 else str(i + 1)
        title = m.group(2).strip() if m.lastindex and m.lastindex >= 2 else ""
        content = text[start:end].strip()
        if content:
            clauses.append(Clause(number=number, title=title, content=content))
    return clauses
